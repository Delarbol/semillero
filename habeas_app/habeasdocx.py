from docx           import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared    import Cm
from docx.shared    import Pt

def negrita(parrafo,texto):
    """Esta función agrega en negrilla el texto en el párrafo dado"""
    parrafo.add_run(texto).bold = True
def crea_habeas(id_solicitante,
                id_afectado,
                fecha,
                title ,
                nom_solicitante,
                ciudad ,
                condi_solici,
                direccion_solicitante,
                email_solicitante ,
                ced_solicitante,
                num_solicitante ,
                nom_afectado ,
                nom_autoridad ,
                fecha_hechos ,
                ced_afectado ,
                num_dias,
                gen_afectado,
                sitio ,
                sujeto_ordeno = None, # <--------------DATOS OPCIONALES
                cargo_txt = None,
                hechos = None):
    # ====creación del documento=======================
    doc = Document()
    # ===================MÁRGENES=========================
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
    # =============FORMATO==========================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
    # =============ENCABEZAD0==================================================
    header = doc.sections[0].header # poner la sección donde irá el encabezado
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r= p.add_run()
    r.add_picture(r'.\static\logo2.png',  width=Cm(18.0))
    # =====================PRIMER PARRAFO======================================
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    negrita(p1,'\nSeñor \nJuez, \n\nE.\tS.\tD.') #NEGRITA
    # =====================SEGUNDO PARRAFO=====================================
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(f'{ciudad}\n{fecha}')
    # =====================TERCER PARRAFO=====================================
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    p3.add_run(f'''Yo, {nom_solicitante} en mi condición de
    {condi_solici}, acudo ante usted, señor juez a fin de solicitarle se
    sirva dar trámite a la petición de habeas corpus en favor de {nom_afectado},
    identificado con {id_afectado} No. {ced_afectado}, con fundamento en lo siguiente:\n''')
    # =====================CUARTO PARRAFO=====================================
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    negrita(p4, 'HECHOS:\n')
    p4.add_run(f'''{nom_afectado} fue aprehendid{gen_afectado} por la
    {nom_autoridad} el pasado {fecha_hechos} por orden de {sujeto_ordeno}.
    Desde entonces, hasta la fecha han transcurrido {num_dias} días
    sin que haya sido indagada o resuelta su situación jurírdica.{nom_afectado}
    se encuentra recluid{gen_afectado} en {sitio}, desde el día {fecha_hechos}
    y el funcionario que ordenó; su aprehensión es {sujeto_ordeno} {cargo_txt}\n''')
    p4.add_run(f'{hechos}')
    # =====================QUINTO PARRAFO=====================================
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    negrita(p5, 'JURAMENTO:\n') #NEGRITA
    p5.add_run('''Bajo la gravedad del juramento, manifiesto que ningún otro
    funcionario conoce o ha decidido sobre esta acción.\n''')
    # =====================SEXTO PARRAFO=====================================
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    negrita(p6, 'FUNDAMENTOS DE DERECHO:\n') #NEGRITA
    p6.add_run('''Esta petición está; fundamentada, señor juez, en los artículos 30
    y 85 de la Constitución Nacional referidos a la privación ilegal de la libertad
    y a la aplicación inmediata de los derechos consagrados en la Constitución,
    Política. Sumado a ello, la Convención Americana de Derechos Humanos establece
    en su artículo 7 el derecho a la libertad personal, ante lo cual, nadie puede ser
    sometido a detención arbitraria ni mucho menos puede ser privado de la libertad,
    salvo por causas y condiciones fijadas por la Constitución y la Ley.
    En el artículo 8 de este instrumento también se plasma como garantía judicial
    el derecho que tiene toda persona a ser oída en un plazo razonable y por la
    autoridad competente para la determinación de sus derechos.\n''')
    p6.add_run('Bien se dispuso por parte de la Corte Constitucional en Sentencia C 187 de 2006 que:\n')
    p6.add_run('''Una interpretación acorde con la Constitución Política supone que,
    después de invocado el habeas corpus, la autoridad judicial encargada de conocer,
    deberá; verificar la existencia de las condiciones que conducen a ordenar que
    el peticionario sea puesto en libertad. Tales condiciones son: i) que la persona está privada de la libertad,
    y ii) que la privación de la libertad o la prolongación de la misma se haya dado
    con violación o quebrantamiento del orden constitucional y legal.
    Una vez demostrado que la privaci&oacute;n de la libertad personal o la
    prolongación de la privaci&oacute;n de la libertad son el resultado de actos
    contrarios a lo dispuesto por el ordenamiento constitucional o legal, la autoridad
    judicial competente deberá; ordenar que la persona sea puesta inmediatamente en libertad\n''')
    # =====================SEPTIMO PARRAFO=====================================
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    negrita(p7, 'SOLICITUD:\n') #NEGRITA
    p7.add_run(f'''Efectuada la verificaciónn de la violación de las garantías constitucionales
    y legales, solicito a usted ordenar la libertad inmediata de {nom_afectado}
    y compulsar copias para que se inicien las investigaciones a que hubiere lugar.\n
    Del señor juez,\n''')
    # =====================OCTAVO PARRAFO=====================================
    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    p8.add_run(f'''{nom_solicitante}\n{id_solicitante}\n {ced_solicitante}\n
    Dirección de notificaciónn electrónica: {email_solicitante}\n
    Dirección de noticación: {direccion_solicitante}\n
    Telefóno: {num_solicitante}''')
    #===============SALVAR DOCUMENTO=============================================
    doc.save('./downloads/habeas_'+nom_solicitante[:4]+ced_solicitante[:-3]+'.docx')

    return
