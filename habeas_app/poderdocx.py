from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt

def negrita(parrafo,texto):
    """Esta función agrega en negrilla el texto en el párrafo dado"""
    parrafo.add_run(texto).bold = True

def creaPoder(tipo_proceso, nom_poder, email, gen_poder, ced_poder, id_poder,
    nom_apo, gen_apo, portadore, email_apo, id_apo, ced_apo,
    nom_contra, id_contra, ced_contra,  gen_contra, email_2,
    num_car = None, email_apo_otro = None):

    # try:
    #     doc = Document('poder_'+nom_apo[:4]+num_car+'.docx')
    #     print('Abrió el mismo')
    # except:
    doc = Document()

# ===================MÁRGENES=========================
    doc.sections[0].left_margin = Cm(2)
    doc.sections[0].top_margin = Cm(2)
    doc.sections[0].bottom_margin = Cm(2)
    doc.sections[0].right_margin = Cm(2)
# =============FORMATO==========================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
# =============ENCABEZAD0=====================================================
    header = doc.sections[0].header # poner la sección donde irá el encabezado
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r= p.add_run()
    r.add_picture(r'.\static\logo2.png',  width=Cm(18.0))
# ===============PRIMER PARRAFO================================================
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run('\nSeñor \nJuez, \n\nE.\tS.\tD.').bold = True
# ================SEGUNDO PARRAFO (REFERENCIA)=================================
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run('\nREFERENCIA:').bold = True
    p2.add_run(f' {tipo_proceso} de ')
    p2.add_run(f'{nom_poder}').bold = True
    p2.add_run(' contra')
    p2.add_run(f' {nom_contra}').bold = True
# ------------------------CUERPO DEL PODER--------------------------------------
# ====================TERCER PARRAFO========================================
    p3 = doc.add_paragraph()
    p3.add_run("\n")
    negrita(p3,f'{nom_poder}')   ### Negrilla
    p3.add_run(", mayor de edad, domiciliad")
    p3.add_run(f"{gen_poder} en la ciudad de Bogotá D.C., identificad{gen_poder} con \
{id_poder} No. {ced_poder} y dirección de notificación electrónica {email}\
, por medio del presente escrito, otorgo ")
    negrita(p3,f"PODER ESPECIAL, AMPLIO Y SUFICIENTE") ### Negrilla
    p3.add_run(f" a ")
    negrita(p3,f"{nom_apo}")  ### Negrilla
    p3.add_run(f", mayor de edad, domiciliad{gen_apo} en Bogotá D.C., \
identificad{gen_apo} con {id_apo}\
 No. {ced_apo}, miembro activo del Consultorio Jurídico de la Universidad \
Externado de Colombia, portador{portadore} del carné No. {num_car}, \
con dirección de notificación electrónica {email_apo}, \
con la finalidad de que en mi nombre y representación, inicie y lleve \
hasta su terminación el ")
    negrita(p3,f'{tipo_proceso} ')###Negrilla
    p3.add_run('contra ')
    negrita(p3,f"{nom_contra} ") ### Negrilla
    p3.add_run(f", mayor de edad, domiciliad{gen_contra} en\u00A0Bogotá\u00A0D.C.,\u00A0\
identificad{gen_contra}\u00A0con\u00A0{id_contra}\u00A0No.\u00A0{ced_contra}\
{email_2} \n\
\nMi apoderad{gen_apo} queda facultad{gen_apo} para solicitar medidas cautelares, \
desistir, renunciar, sustituir, recibir, transigir, asumir el presente poder y \
demás facultades en los términos del artículo 77 del \
Código General del Proceso. \n")

    p4 = doc.add_paragraph()
    p4.add_run(f"Sírvase, Señor Juez, reconocerle personería jurídica a mi apoderad{gen_apo}, \
en los términos y para los efectos del\u00A0presente\u00A0poder. \u00A0\n" )

    p5 = doc.add_paragraph()
    p5.add_run("Del Señor Juez,\n\n\n\n")
    negrita(p5,f"{nom_poder}\n")
    p5.add_run(f"{id_poder} No. {ced_poder}\nAcepto,\n\n\n\n")
    negrita(p5,f"{nom_apo}\n")
    p5.add_run(f"{id_apo} No. {ced_apo} \nCarné Consultorio: {num_car} \n\
Dirección de notificación electrónica: {email_apo}")


    p3.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    p4.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

    #===============SALVAR DOCUMENTO=============================================
    doc.save('./downloads/poder_'+nom_apo[:4]+num_car+'.docx')

    return
